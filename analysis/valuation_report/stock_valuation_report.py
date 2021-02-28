# -*- coding: utf-8 -*-
"""
Created on Sun Nov 19 23:15:49 2017

公司估值框架，与无风险利率结合起来

@author: Dai
"""
ROOT_PATH = "C:/Users/Dai/Desktop/investment"

import datetime
import math
import os

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd

import DaiToolkit as tk

plt.rcParams['font.sans-serif'] = ['SimHei']  # 用来正常显示中文标签

trading_date = datetime.datetime.now().strftime("%Y%m%d")
folder_path = ROOT_PATH + "/股票/个人研究/估值Report/" + trading_date
if not os.path.exists(folder_path):
    os.makedirs(folder_path)

GDP_growth = 0.06


###############################################
#   functions: PE calculation  
###############################################

def rate_implied_PE(riskfree_rate):
    """
    利率隐含的PE水平
    get rate implied PE from riskfree_rate
    rate in (0,+inf)
    """
    return math.log(2) / math.log(1 + riskfree_rate)


def PE_rate_sensitivity(prev_rf, curr_rf):
    """
    利率隐含PE对利率的敏感程度
    get perc chg of PE when rf change
    """
    return (rate_implied_PE(curr_rf) - rate_implied_PE(prev_rf)) / rate_implied_PE(prev_rf)


def sum_CF(growth_rate, rate_implied_PE):
    """
    假设最近一年盈利为1，给定利率PE年数和盈利增长率，判断非折现现金流的和
    get sum of cash flow(not discounted) with given growth rate  
    assume past earning = 1
    growth_rate should in [-1,+inf)
    """
    if growth_rate != 0:
        return (1 + growth_rate) / growth_rate * (math.pow(1 + growth_rate, rate_implied_PE) - 1)
    else:
        return rate_implied_PE


# sum_CF(0.000001,rate_implied_PE)

def implied_growthrate(stock_PE, riskfree_rate):
    """
    给定股票PE和无风险利率，求隐含的相同回本收益下的盈利增长率
    implied growth rate assuming get notional back
    growth rate is average basis
    """
    rate_implied_PE = math.log(2) / math.log(1 + riskfree_rate)
    g1 = -1.0
    g2 = 100.0
    if sum_CF(g2, rate_implied_PE) - stock_PE < 0:
        print("Implied growth rate > 10000% !!!")
        return None
    while abs(g2 - g1) > 0.000001:
        if sum_CF((g1 + g2) / 2, rate_implied_PE) - stock_PE > 0:
            g2 = (g1 + g2) / 2
        elif sum_CF((g1 + g2) / 2, rate_implied_PE) - stock_PE < 0:
            g1 = (g1 + g2) / 2
        else:
            return (g1 + g2) / 2
    return (g1 + g2) / 2


def implied_return(stock_PE, growth_rate, riskfree_rate):
    """
    给定股票PE,无风险利率,盈利增长率，求相同利率PE年限下的隐含回报率
    implied rate of return if we know average growth rate of earnings
    """
    rate_implied_PE = math.log(2) / math.log(1 + riskfree_rate)
    if growth_rate != 0:
        return math.pow(
            ((1 + growth_rate) / growth_rate * (math.pow(1 + growth_rate, rate_implied_PE) - 1)) / stock_PE + 1,
            1 / rate_implied_PE) - 1
    else:
        return math.pow(rate_implied_PE / stock_PE + 1, 1 / rate_implied_PE) - 1


def implied_growth(stock_PE, required_ret, riskfree_rate):
    """
    给定股票PE和目标收益率，求隐含的相同回本收益下的盈利增长率
    """
    # rate_implied_PE = math.log(2)/math.log(1+riskfree_rate)
    g1 = -1.0
    g2 = 100.0
    if implied_return(stock_PE, g2, riskfree_rate) < required_ret:
        print("Implied growth rate > 10000% !!!")
        return None
    while abs(g2 - g1) > 0.000001:
        if implied_return(stock_PE, (g1 + g2) / 2, riskfree_rate) - required_ret > 0:
            g2 = (g1 + g2) / 2
        elif implied_return(stock_PE, (g1 + g2) / 2, riskfree_rate) - required_ret < 0:
            g1 = (g1 + g2) / 2
        else:
            return (g1 + g2) / 2
    return (g1 + g2) / 2


def actual_value(stock_PE, growth_rate, riskfree_rate, tenor=None):
    """
    给定股票PE，无风险利率和盈利增长率，算出股票真实价值，返回PE的形式(假设Earning为1,PE=price)
    """
    if not tenor:
        rate_implied_PE = math.log(2) / math.log(1 + riskfree_rate)
        tenor = rate_implied_PE
    g = growth_rate
    r = riskfree_rate
    if g == r:
        return tenor + stock_PE / math.pow(1 + r, tenor)
    else:
        return (1 + g) / (1 + r) * (1 - math.pow((1 + g) / (1 + r), tenor)) / (
                    1 - (1 + g) / (1 + r)) + stock_PE / math.pow(1 + r, tenor)


###############################################
#   Input 
###############################################
df_sec = pd.read_excel(ROOT_PATH + "/股票/个人研究/估值Report/seclist.xlsx", "Sheet1")
df_sec["SecCode"] = ["0" * (6 - len(str(x))) + str(x) for x in df_sec["SecCode"]]
stock_dict = {x: y for x, y in zip(df_sec["SecCode"], df_sec["SecName"])}
growth_pred_dict = {x: y for x, y in zip(df_sec["SecName"], df_sec["Growth_Pred"])}

# override
# PE_dict ={u"青岛海尔":15.81,u"格力电器":11.48,u"民生银行":4.99,u"中国银行":5.86,
#          u"医药行业":30,u"万科A":8.98,u"金地集团":5.54,
#          u"中证500":22.45,u"上证50":9.78,u"沪深300":11.88,u"创业板指":40.77}

# ------------------------------------------ automatic stock price / pe calculation
# download price
price_dict = {}
for k in list(stock_dict.keys()):
    price_dict[stock_dict[k]] = float(tk.ts.get_realtime_quotes(k)["price"][0])
    print(stock_dict[k] + "下载完毕！[price = " + str(price_dict[stock_dict[k]]) + "]")

# eps dict (need refresh every quarter)
eps_TTM_dict = {}
for k in list(stock_dict.keys()):
    eps_TTM_dict[k] = tk.tusharelocal_get_TTMEPS(k, trading_date)
###if all error / lack data
###get nearest Qtr EPS info of all secs
# tk.tushare_get_nearestqtr_profitdata()

# pe dict    
PE_TTM_dict = {stock_dict[k]: float(price_dict[stock_dict[k]]) / eps_TTM_dict[k] for k in list(stock_dict.keys())}
PE_dict = PE_TTM_dict

## rf rate shibor 1y
shibor_tbl = tk.ts.shibor_data()
shibor_tbl_last = tk.ts.shibor_data(int(datetime.datetime.now().strftime("%Y")) - 1)
shibor_tbl.index = shibor_tbl["date"]
shibor_tbl_last.index = shibor_tbl_last["date"]
del shibor_tbl["date"]
del shibor_tbl_last["date"]
riskfree_rate = shibor_tbl["1Y"][-1] / 100.0
# plot shibor curve
plt.figure(figsize=(12, 6))
plt.plot(shibor_tbl.iloc[-1, :].values, label="最新: " + str(shibor_tbl.index[-1])[:10], lw=5, ls="-.")
plt.plot(shibor_tbl.iloc[0, :].values, label="年初: " + str(shibor_tbl.index[0])[:10])
plt.plot(shibor_tbl_last.iloc[len(shibor_tbl_last) / 2, :].values,
         label="去年年中: " + str(shibor_tbl_last.index[len(shibor_tbl_last) / 2])[:10])
plt.plot(shibor_tbl_last.iloc[0, :].values, label="去年年初: " + str(shibor_tbl_last.index[0])[:10])
plt.xticks(list(range(len(shibor_tbl.columns))), shibor_tbl.columns)
plt.grid(linestyle='--', alpha=0.7)
plt.title(trading_date + " SHIBOR curve")
plt.legend()
plt.savefig(folder_path + "/SHIBOR_curve.png", dpi=100, bbox_inches='tight')

# ----------------------- rf = 1Y shibor
# riskfree_rate = 0.035
# GDP_growth = 0.06
print("日期: " + str(datetime.date.today()))
print("无风险利率(shibor 1y): " + str(riskfree_rate))
print("无风险利率隐含PE: " + str(rate_implied_PE(riskfree_rate)))
print("GDP名义增长率预测: " + str(GDP_growth))

summary_tbl = {}
for k, v in list(PE_dict.items()):
    summary_tbl[k] = {}
    summary_tbl[k]["PE"] = v
    summary_tbl[k]["implied_growth"] = implied_growthrate(v, riskfree_rate)
    summary_tbl[k]["10%ret_implied_growth"] = implied_growth(v, 0.10, riskfree_rate)
    summary_tbl[k]["0_growth_implied_return"] = implied_return(v, 0.0, riskfree_rate)
    summary_tbl[k]["GDP_growth_implied_return"] = implied_return(v, GDP_growth, riskfree_rate)
    summary_tbl[k]["pred_growth"] = growth_pred_dict[k]
    summary_tbl[k]["pred_growth_implied_return"] = implied_return(v, growth_pred_dict[k], riskfree_rate)

summary_tbl = pd.DataFrame(summary_tbl).T
summary_tbl = summary_tbl[
    ["PE", "implied_growth", "10%ret_implied_growth", "0_growth_implied_return", "GDP_growth_implied_return",
     "pred_growth", "pred_growth_implied_return"]]
print(summary_tbl)

summary_value_tbl = {}
for k, v in list(PE_dict.items()):
    summary_value_tbl[k] = {}
    summary_value_tbl[k]["PE"] = v
    summary_value_tbl[k]["pred_growth"] = growth_pred_dict[k]
    summary_value_tbl[k]["3y_value_discount"] = v / actual_value(v, growth_pred_dict[k], riskfree_rate, tenor=3) - 1
    summary_value_tbl[k]["5y_value_discount"] = v / actual_value(v, growth_pred_dict[k], riskfree_rate, tenor=5) - 1
    summary_value_tbl[k]["7y_value_discount"] = v / actual_value(v, growth_pred_dict[k], riskfree_rate, tenor=7) - 1
    summary_value_tbl[k]["10y_value_discount"] = v / actual_value(v, growth_pred_dict[k], riskfree_rate, tenor=10) - 1
    summary_value_tbl[k]["PE_value_discount"] = v / actual_value(v, growth_pred_dict[k], riskfree_rate) - 1

summary_value_tbl = pd.DataFrame(summary_value_tbl).T
summary_value_tbl = summary_value_tbl[
    ["PE", "pred_growth", "3y_value_discount", "5y_value_discount", "7y_value_discount", "10y_value_discount",
     "PE_value_discount"]]
print(summary_value_tbl)

# 隐含收益率和增速关系
plt.figure(figsize=(12, 6))
for c in summary_value_tbl.index[:]:
    plt.plot([3, 5, 7, 10, rate_implied_PE(riskfree_rate)], summary_value_tbl.T.loc[:, c][2:], marker="o", label=c)
plt.xticks([3, 5, 7, 10, rate_implied_PE(riskfree_rate)], [3, 5, 7, 10, rate_implied_PE(riskfree_rate)])
plt.legend()
plt.axhline(y=0, linestyle="--", color="darkblue")
plt.grid(linestyle='--', alpha=0.7)
plt.title(trading_date + "（预计盈利增速下）年限和内在价值折价关系")
plt.savefig(folder_path + "/undervalue_growth.png", dpi=100, bbox_inches='tight')

# 隐含收益率和增速关系
plt.figure(figsize=(12, 6))
g_list = [round(x, 5) for x in np.arange(-0.10, 0.4, 0.01)]
for k, v in list(PE_dict.items()):
    plt.plot(g_list, [implied_return(v, g, riskfree_rate) for g in g_list], label=k)
    plt.scatter(growth_pred_dict[k], implied_return(v, growth_pred_dict[k], riskfree_rate))
plt.axvline(x=0, linestyle="--")
plt.axvline(x=GDP_growth, linestyle="--")
plt.axhline(y=riskfree_rate, linestyle="--", color="darkblue")
plt.text(0.35, riskfree_rate + 0.01, str(round(riskfree_rate * 100, 2)) + "% 1y-Rf Rate", verticalalignment='center')
plt.text(0 + 0.003, 0.3, "0% Growth Rate", rotation=90, verticalalignment='center')
plt.text(GDP_growth + 0.003, 0.29, str(GDP_growth * 100) + "% GDP Growth Rate", rotation=90, verticalalignment='center')
plt.xticks(np.arange(-0.05, 0.41, 0.05))
plt.yticks(np.arange(0, 0.35, 0.02))
plt.xlabel("Growth Rate")
plt.ylabel("implied return")
plt.legend()
plt.grid(linestyle='--', alpha=0.7)
plt.title(trading_date + " 收益率和盈利增速(时间跨度" + str(round(rate_implied_PE(riskfree_rate), 2)) + "年)")
plt.savefig(folder_path + "/impret_growth.png", dpi=100, bbox_inches='tight')

# 增长率一定下，PE和隐含收益率关系图
plt.figure(figsize=(12, 6))
for g in range(0, 16, 3):
    g = g / 100.0
    plt.plot(np.arange(8, 30, 0.5), [implied_return(pe, g, riskfree_rate) for pe in np.arange(8, 30, 0.5)],
             label=str(g * 100) + "%")
plt.grid(linestyle='--', alpha=0.7)
plt.title(trading_date + " 增长率一定下, PE和隐含收益率关系图, 无风险利率= " + str(riskfree_rate * 100) + "% (时间跨度" + str(
    round(rate_implied_PE(riskfree_rate), 2)) + "年)")
plt.legend()
plt.savefig(folder_path + "/PEimpret_growth.png", dpi=100, bbox_inches='tight')

# 现金选择权剩余时间，在剩余时间内下跌至rf对应PE盈亏平衡
drop_wait_dict = {}
for tkr_name in list(PE_dict.keys()):
    ret_small = implied_return(PE_dict[tkr_name], growth_pred_dict[tkr_name], riskfree_rate)
    ret_big = implied_return(rate_implied_PE(riskfree_rate), growth_pred_dict[tkr_name], riskfree_rate)
    drop_wait_period = (rate_implied_PE(riskfree_rate) * math.log((1 + ret_small) / (1 + ret_big))) / (
        math.log((1 + riskfree_rate) / (1 + ret_big)))
    if drop_wait_period > 0:
        drop_wait_dict[tkr_name] = {"option_period(curr-rf)": drop_wait_period, "curr_growth_pred_ret": ret_small,
                                    "rf_PE_pred_ret": ret_big, "curr_PE": PE_dict[tkr_name],
                                    "rf_PE": rate_implied_PE(riskfree_rate)}
    else:
        drop_wait_dict[tkr_name] = {"option_period(curr-rf)": float("nan"), "curr_growth_pred_ret": ret_small,
                                    "rf_PE_pred_ret": ret_big, "curr_PE": PE_dict[tkr_name],
                                    "rf_PE": rate_implied_PE(riskfree_rate)}

drop_wait_dict = pd.DataFrame(drop_wait_dict).T
print(drop_wait_dict)

#############################################################
# Docx Report Generate
from docx import Document
from docx.shared import Inches

doc_path = folder_path + "/Evaluation_Report.docx"


def docx_add_table(df, index=True):
    """add table to the docx"""
    table = document.add_table(rows=len(df.index) + 1, cols=len(df.columns) + 1, style='Table Grid')
    for numc, c in enumerate(df.columns):
        table.rows[0].cells[numc + 1].text = " ".join(c.split("_"))

    for i in range(len(df.index)):
        table.rows[i + 1].cells[0].text = df.index[i]
        table.rows[i + 1].cells[0].width = Inches(1.0)
        for numc, c in enumerate(df.columns):
            table.rows[i + 1].cells[numc + 1].text = str(round(df[c][i], 2))
            table.rows[i + 1].cells[numc + 1].width = Inches(0.8)

        # 新建一个文档


document = Document()
document.add_heading(trading_date + ' 股票池估值报告', level=0)
document.add_paragraph("日期: " + str(datetime.date.today()))
document.add_paragraph("无风险利率(shibor 1y): " + str(riskfree_rate))
document.add_paragraph("无风险利率隐含PE: " + str(rate_implied_PE(riskfree_rate)))
document.add_paragraph("GDP名义增长率预测: " + str(GDP_growth))
document.add_paragraph("")

document.add_heading('一、SHIBOR Curve', level=1)
document.add_picture(folder_path + "/SHIBOR_curve.png", width=Inches(6.0))

# 插入表格
document.add_heading('二、隐含收益率分析表', level=1)
docx_add_table(summary_tbl, index=True)
document.add_paragraph("")

document.add_heading('三、隐含收益率和增速关系', level=1)
document.add_picture(folder_path + "/impret_growth.png", width=Inches(6.0))
document.add_paragraph("")

document.add_heading('四、折价分析表', level=1)
docx_add_table(summary_value_tbl)
document.add_paragraph("")

document.add_heading('五、预计盈利增速下, 年限和内在价值折价关系', level=1)
document.add_picture(folder_path + "/undervalue_growth.png", width=Inches(6.0))
document.add_paragraph("")

document.add_heading('六、增长率一定下，PE和隐含收益率关系图', level=1)
document.add_picture(folder_path + "/PEimpret_growth.png", width=Inches(6.0))
document.add_paragraph("")

document.add_heading('七、现金选择权剩余时间，在剩余时间内下跌至rf对应PE盈亏平衡', level=1)
docx_add_table(drop_wait_dict)
document.add_paragraph("")

document.add_page_break()
document.save(doc_path)
