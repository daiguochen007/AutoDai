# -*- coding: utf-8 -*-
"""
Created on Sat Jan 05 14:40:25 2019

@author: Dai
"""
# %%
ROOT_PATH = "C:/Users/Dai/Desktop/investment"
import os
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib as mpl

import DaiToolkit as tk


mpl.rcParams.update(mpl.rcParamsDefault)
plt.rcParams['font.sans-serif'] = ['SimHei']  # 用来正常显示中文标签

df_sec = pd.read_excel(ROOT_PATH + "/股票/个人研究/Rebalancing/sec_pool.xlsx", "Sheet1")
df_sec["SecCode"] = ["0" * (6 - len(str(x))) + str(x) for x in df_sec["SecCode"]]

group_list = list(set(df_sec["SecGroup"]))
selected_group = ["保险"]  # ,u"农商行",u"城商行",u"国有大行",u"股份制银行",u"房地产",u"白色家电",u"上证50"
df_seclist = df_sec.loc[df_sec["SecGroup"].isin(selected_group), ["SecCode"]]
sec_list = list(df_seclist["SecCode"].values)
dict_localallsec = tk.tusharelocal_get_allsecnames()
for sec in sec_list:
    if sec not in list(dict_localallsec.keys()):
        dict_localallsec[sec] = sec
df_seclist["SecName"] = [dict_localallsec[x] for x in df_seclist["SecCode"]]

folder_path = ROOT_PATH + "/股票/个人研究/Rebalancing/" + "_".join(selected_group)
if not os.path.exists(folder_path):
    os.makedirs(folder_path)

# %%
######################################## download
for seccode in sec_list:
    try:
        # 1/0
        df = tk.tusharelocal_get_history(seccode)
    except:
        print("[" + seccode + "] not in local. Downloading...")
        tk.tushare_get_history(seccode)

# %%
########################################
df_price = pd.DataFrame()
df_PE = pd.DataFrame()

for seccode in sec_list:
    pass
    df = tk.tusharelocal_get_history(seccode)
    df.index = pd.DatetimeIndex(list(map(str, df["trade_date"])))
    temp = df[["close"]]
    temp.columns = [seccode]
    df_price = pd.concat([df_price, temp], axis=1)
    temp = df[["PE_TTM"]]
    temp.columns = [seccode]
    df_PE = pd.concat([df_PE, temp], axis=1)
    print(seccode + " finished!")

df_price = df_price.fillna(method="ffill")
df_PE = df_PE.fillna(method="ffill")

start_date = max(df_price.dropna(how="any").index[0], df_PE.dropna(how="any").index[0]).strftime("%Y%m%d")
end_date = df_price.dropna(how="any").index[-1].strftime("%Y%m%d")

df_PE = df_PE[df_PE.index >= start_date]
df_price = df_price[df_price.index >= start_date]
df_ret = df_price / df_price.shift(1) - 1
df_ret = df_ret.fillna(0)

df_perf_stats = pd.DataFrame()
######################################## equal weight
df_weight = pd.DataFrame(index=df_ret.index, columns=df_ret.columns)
for idx in df_weight.index:
    df_weight.loc[idx] = 1.0 / len(sec_list)
ret_equal_port = (df_ret * df_weight).sum(axis=1)
perf_equal_port = (ret_equal_port + 1).cumprod()
perf_allsecs = (df_ret + 1).cumprod()
print("equal weight finished.")

######################################## static qty with equal weight at beginning
start_qty_list = 1.0 / df_price.iloc[0, :]
perf_staticqty_port = (start_qty_list * df_price).sum(axis=1)
perf_staticqty_port = perf_staticqty_port / perf_staticqty_port[0]
ret_staticqty_port = perf_staticqty_port / perf_staticqty_port.shift(1) - 1
perf_stats = pd.DataFrame(tk.perf_stats(ret_staticqty_port, factor_returns=ret_staticqty_port), columns=["static qty"])
df_perf_stats = pd.concat([df_perf_stats, perf_stats], axis=1)
print("static qty finished.")
# equal weight stat
perf_stats = pd.DataFrame(tk.perf_stats(ret_equal_port, factor_returns=ret_staticqty_port), columns=["equal weight"])
df_perf_stats = pd.concat([df_perf_stats, perf_stats], axis=1)

######################################## weight by rule 1/PE
df_weight = pd.DataFrame(index=df_ret.index, columns=df_ret.columns)

for idx in df_weight.index:
    temp = 1.0 / (df_PE.iloc[list(df_PE.index).index(idx) - 1])
    # if PE>20  or neg PE set 0 weight
    if sum([x if (x >= 0.05 and x >= 0) else 0 for x in temp]) == 0:
        temp = temp.values / sum(temp)
    else:
        temp = [x if (x >= 0.05 and x >= 0) else 0 for x in temp]
        temp = np.array(temp) / sum(temp)
    df_weight.loc[idx] = temp

ret_PE_port = (df_ret * df_weight).sum(axis=1)
perf_stats = pd.DataFrame(tk.perf_stats(ret_PE_port, factor_returns=ret_staticqty_port), columns=["1/PE weight"])
df_perf_stats = pd.concat([df_perf_stats, perf_stats], axis=1)
perf_PE_port = (ret_PE_port + 1).cumprod()
print("weight by rule 1/PE finished.")


######################################## weight by rule relative ret
def relative_ret_rebalancing(df_ret, weight_lower=0, weight_upper=0.5, weight_move_interval=50, relative_perf=0.03):
    """
    一对股票基于ret变化rebalancing
    """
    sec_list = list(df_ret.columns)
    df_weight = pd.DataFrame(index=df_ret.index, columns=df_ret.columns)
    weight_move = (weight_upper - weight_lower) / weight_move_interval
    df_weight.iloc[0, :] = 1.0 / len(sec_list)
    # df_weight.iloc[0,:] = 1.0/(df_PE.iloc[list(df_PE.index).index(df_weight.index[0])-1])
    # df_weight.iloc[0,:] = df_weight.iloc[0,:]/sum(df_weight.iloc[0,:])
    for idx in df_weight.index[1:]:
        df_weight.loc[idx] = df_weight.iloc[list(df_weight.index).index(idx) - 1]
        ret_list = df_ret.iloc[list(df_ret.index).index(idx) - 1]
        if max(ret_list) - min(ret_list) > relative_perf:
            max_ret_ind = ret_list.idxmax()
            min_ret_ind = ret_list.idxmin()
            if df_weight.loc[idx, max_ret_ind] < weight_lower + weight_move or df_weight.loc[
                idx, min_ret_ind] > weight_upper - weight_move:
                continue
            else:
                df_weight.loc[idx, max_ret_ind] = df_weight.loc[idx, max_ret_ind] - weight_move
                df_weight.loc[idx, min_ret_ind] = df_weight.loc[idx, min_ret_ind] + weight_move

    ret_relative_port = (df_ret * df_weight).sum(axis=1)
    perf_stats = pd.DataFrame(tk.perf_stats(ret_relative_port, factor_returns=ret_staticqty_port),
                              columns=["Relative Chg " + str(int(relative_perf * 100)) + "%"])
    perf_relative_port_perc = (ret_relative_port + 1).cumprod()
    return perf_relative_port_perc, perf_stats, df_weight


perf_relative_port_5perc, perf_stats, df_weight = relative_ret_rebalancing(df_ret, relative_perf=0.05)
df_perf_stats = pd.concat([df_perf_stats, perf_stats], axis=1)
print("weight by rule 5% relative ret finished.")

perf_relative_port_3perc, perf_stats, df_weight = relative_ret_rebalancing(df_ret, relative_perf=0.03)
df_perf_stats = pd.concat([df_perf_stats, perf_stats], axis=1)
print("weight by rule 3% relative ret finished.")

######################################## weight by rule relative ret
perf_relative_port_2perc, perf_stats, df_weight = relative_ret_rebalancing(df_ret, relative_perf=0.02)
df_perf_stats = pd.concat([df_perf_stats, perf_stats], axis=1)
print("weight by rule 2% relative ret finished.")

######################################## weight by rule relative ret
perf_relative_port_1perc, perf_stats, df_weight = relative_ret_rebalancing(df_ret, relative_perf=0.01)
df_perf_stats = pd.concat([df_perf_stats, perf_stats], axis=1)
print("weight by rule 1% relative ret finished.")

# %%
################### plot final result
plt.figure(figsize=(12, 10))
for sec in perf_allsecs.columns:
    plt.plot(perf_allsecs[sec], alpha=0.3, label=dict_localallsec[sec])
plt.plot(perf_equal_port, label="Equal Weighted Portfolio")
plt.plot(perf_staticqty_port, label="Static Qty Portfolio")
plt.plot(perf_PE_port, label="1/PE Weight Portfolio")
plt.plot(perf_relative_port_5perc, label="Relative Weight Portfolio (5% Trigger)")
plt.plot(perf_relative_port_3perc, label="Relative Weight Portfolio (3% Trigger)")
plt.plot(perf_relative_port_2perc, label="Relative Weight Portfolio (2% Trigger)")
plt.plot(perf_relative_port_1perc, label="Relative Weight Portfolio (1% Trigger)")
plt.grid(ls="--", alpha=0.5)
plt.title("|".join(selected_group) + " Rebalance Backtest " + start_date + "-" + end_date)
plt.legend(loc='upper center', bbox_to_anchor=(0.5, -0.1), ncol=4, fontsize=12)
plt.savefig(folder_path + "/backtest_perf.png", dpi=100, bbox_inches='tight')
plt.show()

# plot final result
plt.figure(figsize=(12, 10))
plt.plot(perf_equal_port - perf_staticqty_port, label="Equal weight - Static Qty")
plt.plot(perf_PE_port - perf_staticqty_port, label="1/PE weight - Static Qty")
plt.plot(perf_relative_port_5perc - perf_staticqty_port, label="Relative weight (5% Trigger) - Static Qty")
plt.plot(perf_relative_port_3perc - perf_staticqty_port, label="Relative weight (3% Trigger) - Static Qty")
plt.plot(perf_relative_port_2perc - perf_staticqty_port, label="Relative weight (2% Trigger) - Static Qty")
plt.plot(perf_relative_port_1perc - perf_staticqty_port, label="Relative weight (1% Trigger) - Static Qty")
plt.grid(ls="--", alpha=0.5)
plt.title("|".join(selected_group) + " Rebalance Backtest Relative Perf " + start_date + "-" + end_date)
plt.legend(fontsize=12)
plt.savefig(folder_path + "/backtest_relperf.png", dpi=100, bbox_inches='tight')
plt.show()

# plot weight
plt.figure(figsize=(12, 3))
plt.stackplot(df_weight.index, list(map(list, df_weight.T.values)), labels=[dict_localallsec[x] for x in df_weight.columns])
plt.title("|".join(selected_group) + " Relative 1% weight - Weight Change (area) " + start_date + "-" + end_date)
plt.legend(loc='upper center', bbox_to_anchor=(0.5, -0.1), ncol=4, fontsize=12)
plt.savefig(folder_path + "/weight_chg_area.png", dpi=100, bbox_inches='tight')
plt.show()

plt.figure(figsize=(12, 3))
for c in df_weight.columns:
    plt.plot(df_weight[c], label=dict_localallsec[c])
plt.title("|".join(selected_group) + " Relative 1% weight - Weight Change (line) " + start_date + "-" + end_date)
plt.legend(loc='upper center', bbox_to_anchor=(0.5, -0.1), ncol=4, fontsize=12)
plt.grid(ls="--", alpha=0.5)
plt.savefig(folder_path + "/weight_chg_line.png", dpi=100, bbox_inches='tight')
plt.show()

# pe
plt.figure(figsize=(12, 5))
for c in df_PE.columns:
    plt.plot(df_PE[c], label=dict_localallsec[c])
plt.ylim(0, 30)
plt.title("|".join(selected_group) + " PE " + start_date + "-" + end_date)
plt.legend(loc='upper center', bbox_to_anchor=(0.5, -0.1), ncol=4, fontsize=12)
plt.grid(ls="--", alpha=0.5)
plt.savefig(folder_path + "/pe.png", dpi=100, bbox_inches='tight')
plt.show()

# %%
#############################################################
# Report Generate
from docx import Document
from docx.shared import Inches, Pt

doc_path = folder_path + "/Rebalancing_Report.docx"


def docx_add_table(df, index=True):
    """add table to the docx"""
    if index:
        table = document.add_table(rows=len(df.index) + 1, cols=len(df.columns) + 1, style='Table Grid')
    else:
        table = document.add_table(rows=len(df.index) + 1, cols=len(df.columns), style='Table Grid')
    for numc, c in enumerate(df.columns):
        if not index:
            numc -= 1
        table.rows[0].cells[numc + 1].text = " ".join(c.split("_"))
        table.rows[0].cells[numc + 1].paragraphs[0].runs[0].font.size = Pt(11)
        table.rows[0].cells[numc + 1].paragraphs[0].runs[0].font.bold = True
    for i in range(len(df.index)):
        if index:
            table.rows[i + 1].cells[0].text = df.index[i]
            table.rows[i + 1].cells[0].width = Inches(2.0)
            table.rows[i + 1].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
        for numc, c in enumerate(df.columns):
            if not index:
                numc -= 1
            x = df[c].iloc[i]
            table.rows[i + 1].cells[numc + 1].text = str(round(x, 4)) if (type(x) != str and type(x) != str) else x
            table.rows[i + 1].cells[numc + 1].width = Inches(0.35)
            table.rows[i + 1].cells[numc + 1].paragraphs[0].runs[0].font.size = Pt(11)


# 新建一个文档
document = Document()
document.add_heading("|".join(selected_group) + ' Rebalance Backtest', level=0)
document.add_paragraph("开始日期: " + start_date)
document.add_paragraph("结束日期: " + end_date)

document.add_heading('一、证券池', level=1)
docx_add_table(df_seclist, index=False)

document.add_heading('二、Rebalancing回测分析', level=1)
docx_add_table(df_perf_stats, index=True)

document.add_heading('三、Rebalancing回测净值表现', level=1)
document.add_picture(folder_path + "/backtest_perf.png", width=Inches(6.0))

document.add_heading('四、Rebalancing回测净值相对表现', level=1)
document.add_picture(folder_path + "/backtest_relperf.png", width=Inches(6.0))

document.add_heading('五、1%相对波动策略仓位变动（area）', level=1)
document.add_picture(folder_path + "/weight_chg_area.png", width=Inches(6.0))

document.add_heading('六、1%相对波动策略仓位变动（line）', level=1)
document.add_picture(folder_path + "/weight_chg_line.png", width=Inches(6.0))

document.add_heading('七、历史PE', level=1)
document.add_picture(folder_path + "/pe.png", width=Inches(6.0))

document.add_page_break()
document.save(doc_path)
